import json
import re
from datetime import datetime, timedelta

from fastapi import APIRouter, Depends, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import case, delete, func, select
from sqlalchemy.orm.attributes import flag_modified

from . import ai, exports, segment
from .auth import (
    FEATURES, allowed_bot_ids, can_delete, current_user, ensure_bot_access,
    ensure_funnel_access, hash_password, make_token, normalize_permissions,
    parse_token, require, require_auth, require_delete, require_owner,
    user_can_bot, user_funnel_ids, user_permissions, verify_password,
)
from .db import SessionLocal
from .graph import GraphError, compile_graph
from .models import (
    AIRequest,
    Bot,
    Broadcast,
    BroadcastRecipient,
    ButtonClick,
    Funnel,
    FunnelBot,
    FunnelRun,
    Message,
    NodeVisit,
    ScheduledJob,
    Subscriber,
    SubscriberTag,
    Tag,
    User,
)

from .logging_setup import event_logger

log = event_logger()

router = APIRouter(prefix="/api")


async def get_session():
    async with SessionLocal() as session:
        yield session
        await session.commit()


# ---------- health ----------

@router.get("/health")
async def health():
    """Проверка живости для деплоя и мониторинга (без авторизации)."""
    from .bot.runner import manager

    expected = 0
    try:
        async with SessionLocal() as s:
            expected = (await s.execute(
                select(func.count(Bot.id)).where(Bot.is_active == True)  # noqa: E712
            )).scalar() or 0
        db_ok = True
    except Exception:  # noqa: BLE001
        db_ok = False

    running = len(manager.bots)
    # «ok» только когда включённые боты действительно поллятся. Раньше здесь
    # хватало живого HTTP, и деплой рапортовал об успехе, даже если все боты
    # лежали и /start ничего не делал.
    healthy = db_ok and running >= expected
    return {
        "status": "ok" if healthy else "degraded",
        "db": db_ok,
        "bots_running": running,
        "bots_expected": expected,
        "time": datetime.utcnow().isoformat(),
    }


# ---------- логи (только владелец) ----------

@router.get("/logs", dependencies=[Depends(require("logs", "view"))])
async def get_logs(lines: int = 200, only_errors: bool = False, q: str = ""):
    """Последние строки лога — чтобы смотреть из админки, а не по SSH."""
    from pathlib import Path

    from .config import settings

    fname = "sendbot-error.log" if only_errors else "sendbot.log"
    path = Path(settings.log_dir) / fname
    if not path.exists():
        return {"lines": [], "file": str(path), "size_kb": 0,
                "note": "Файл ещё не создан — логи появятся после первых событий"}

    lines = max(10, min(lines, 2000))
    # читаем хвост файла, не загружая его целиком
    with open(path, "rb") as f:
        f.seek(0, 2)
        size = f.tell()
        chunk = min(size, 512 * 1024)
        f.seek(size - chunk)
        text = f.read().decode("utf-8", errors="replace")
    rows = [r for r in text.splitlines() if r.strip()]
    if q:
        needle = q.lower()
        rows = [r for r in rows if needle in r.lower()]
    return {
        "lines": rows[-lines:],
        "file": str(path),
        "size_kb": round(size / 1024, 1),
        "files": sorted(p.name for p in Path(settings.log_dir).glob("sendbot*.log*")),
    }


@router.websocket("/logs/ws")
async def logs_ws(ws):
    """Стриминг логов по сокету: сразу отдаём хвост файла, дальше — только
    новые строки по мере появления. Никакого периодического опроса."""
    import asyncio
    from pathlib import Path

    from .auth import has_perm
    from .config import settings

    await ws.accept()
    # токен передаём параметром: браузер не умеет слать заголовки в WebSocket
    data = parse_token(ws.query_params.get("token", ""))
    ws_user = None
    if data:
        async with SessionLocal() as s_:
            ws_user = await s_.get(User, data["uid"])
    if not ws_user or not ws_user.is_active or not has_perm(ws_user, "logs", "view"):
        await ws.send_json({"error": "Нет прав на просмотр логов"})
        await ws.close()
        return

    only_errors = ws.query_params.get("only_errors") == "true"
    path = Path(settings.log_dir) / ("sendbot-error.log" if only_errors else "sendbot.log")
    try:
        while not path.exists():          # файл появится с первым событием
            await asyncio.sleep(1)
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            f.seek(0, 2)
            size = f.tell()
            f.seek(max(0, size - 64 * 1024))
            tail = f.read().splitlines()[-200:]
            if tail:
                await ws.send_json({"lines": tail, "initial": True})
            inode = path.stat().st_ino
            while True:
                line = f.readline()
                if line:
                    await ws.send_json({"lines": [line.rstrip()]})
                    continue
                await asyncio.sleep(0.4)
                # файл мог смениться при ротации — переоткрываем
                try:
                    if path.stat().st_ino != inode:
                        break
                except FileNotFoundError:
                    break
    except Exception:  # noqa: BLE001 — клиент отключился
        pass
    finally:
        try:
            await ws.close()
        except Exception:  # noqa: BLE001
            pass


# ---------- auth ----------

class LoginIn(BaseModel):
    login: str = "admin"
    password: str


async def _limit_to_my_bots(user, session, column):
    """Условие «только мои боты» для любого запроса, или None если ограничений нет."""
    ids = await allowed_bot_ids(user, session)
    return None if ids is None else column.in_(ids or [-1])


async def _visible_funnel_ids(user, session) -> list[int] | None:
    """Какие воронки пользователь вправе видеть: пересечение «его ботов»
    и явно выданного списка воронок. None = все."""
    bots = await allowed_bot_ids(user, session)
    by_bot = None
    if bots is not None:
        rows = (await session.execute(
            select(FunnelBot.funnel_id).where(FunnelBot.bot_id.in_(bots or [-1]))
        )).all()
        by_bot = {r[0] for r in rows}
    explicit = user_funnel_ids(user)
    if by_bot is None and explicit is None:
        return None
    if by_bot is None:
        return list(explicit)
    if explicit is None:
        return list(by_bot)
    return list(by_bot & set(explicit))


async def _my_funnel(user, session, funnel_id: int):
    """Воронка с проверкой доступа."""
    f = await session.get(Funnel, funnel_id)
    if f is None:
        raise HTTPException(404, "Воронка не найдена")
    ids = await _visible_funnel_ids(user, session)
    if ids is not None and funnel_id not in ids:
        raise HTTPException(403, "Нет доступа к этой воронке")
    return f


async def _my_subscriber(user, session, sub_id: int):
    """Подписчик с проверкой, что он принадлежит доступному боту."""
    sub = await session.get(Subscriber, sub_id)
    if sub is None:
        raise HTTPException(404, "Подписчик не найден")
    if not user_can_bot(user, sub.bot_id):
        raise HTTPException(403, "Нет доступа к этому подписчику")
    return sub


def _user_public(user) -> dict:
    """Данные о себе для интерфейса: и при входе, и при обновлении страницы."""
    return {
        "id": user.id,
        "login": user.login,
        "name": user.name,
        "role": user.role,
        "bot_ids": user.bot_ids or [],
        "funnel_ids": (user.funnel_ids or []) if user.role != "owner" else [],
        "permissions": user_permissions(user),
    }


@router.post("/auth/login")
async def login(body: LoginIn, session=Depends(get_session)):
    res = await session.execute(select(User).where(User.login == body.login.strip()))
    user = res.scalar_one_or_none()
    if user is None or not user.is_active or not verify_password(body.password, user.password_hash):
        log.warning("Неудачный вход: логин «%s»", body.login.strip())
        raise HTTPException(401, "Неверный логин или пароль")
    user.last_login_at = datetime.utcnow()
    log.info("Вход в админку: %s (%s)", user.login, user.role)
    return {
        "token": make_token(user),
        # состав должен совпадать с /auth/me — интерфейс сразу после входа
        # берёт права отсюда, а не из отдельного запроса
        "user": _user_public(user),
    }


@router.get("/auth/me", dependencies=[Depends(require_auth)])
async def auth_me(user=Depends(current_user)):
    return _user_public(user)


@router.get("/permissions/features", dependencies=[Depends(require_owner)])
async def permission_features():
    """Справочник разделов для экрана «Команда»."""
    return FEATURES


class ChangePwIn(BaseModel):
    old_password: str
    new_password: str


@router.post("/auth/password")
async def change_password(body: ChangePwIn, user=Depends(current_user), session=Depends(get_session)):
    db_user = await session.get(User, user.id)
    if not verify_password(body.old_password, db_user.password_hash):
        raise HTTPException(400, "Текущий пароль неверный")
    if len(body.new_password) < 6:
        raise HTTPException(400, "Пароль слишком короткий (минимум 6 символов)")
    db_user.password_hash = hash_password(body.new_password)
    return {"ok": True}


# ---------- команда (только владелец) ----------

class UserIn(BaseModel):
    login: str
    name: str = ""
    password: str | None = None
    role: str = "staff"
    bot_ids: list[int] = []          # пусто = все боты
    funnel_ids: list[int] = []       # пусто = все воронки доступных ботов
    permissions: dict = {}           # {"funnels": "edit", ...}; у владельца не важно
    is_active: bool = True


@router.get("/users", dependencies=[Depends(require_owner)])
async def list_users(session=Depends(get_session)):
    users = (await session.execute(select(User).order_by(User.created_at))).scalars().all()
    return [
        {"id": u.id, "login": u.login, "name": u.name, "role": u.role,
         "is_active": u.is_active, "bot_ids": u.bot_ids or [],
         "funnel_ids": u.funnel_ids or [],
         "permissions": user_permissions(u),
         "last_login_at": u.last_login_at.isoformat() if u.last_login_at else None}
        for u in users
    ]


@router.post("/users", dependencies=[Depends(require_owner)])
async def create_user(body: UserIn, session=Depends(get_session)):
    login = body.login.strip()
    if not login or not body.password or len(body.password) < 6:
        raise HTTPException(400, "Нужен логин и пароль от 6 символов")
    exists = (await session.execute(select(User).where(User.login == login))).scalar_one_or_none()
    if exists:
        raise HTTPException(400, "Такой логин уже есть")
    u = User(login=login, name=body.name or login,
             password_hash=hash_password(body.password),
             role="owner" if body.role == "owner" else "staff",
             bot_ids=body.bot_ids or [], funnel_ids=body.funnel_ids or [],
             is_active=body.is_active,
             permissions=normalize_permissions(body.permissions))
    session.add(u)
    await session.flush()
    log.info("Создан пользователь: %s (роль %s, боты %s)", u.login, u.role, u.bot_ids or "все")
    return {"id": u.id}


@router.put("/users/{user_id}", dependencies=[Depends(require_owner)])
async def update_user(user_id: int, body: UserIn, session=Depends(get_session)):
    u = await session.get(User, user_id)
    if u is None:
        raise HTTPException(404, "Пользователь не найден")
    u.name = body.name or u.name
    u.role = "owner" if body.role == "owner" else "staff"
    u.bot_ids = body.bot_ids or []
    u.funnel_ids = body.funnel_ids or []
    u.is_active = body.is_active
    u.permissions = normalize_permissions(body.permissions)
    if body.password:
        if len(body.password) < 6:
            raise HTTPException(400, "Пароль слишком короткий")
        u.password_hash = hash_password(body.password)
    return {"ok": True}


@router.delete("/users/{user_id}", dependencies=[Depends(require_owner)])
async def delete_user(user_id: int, user=Depends(current_user), session=Depends(get_session)):
    if user_id == user.id:
        raise HTTPException(400, "Нельзя удалить самого себя")
    owners = (await session.execute(select(func.count(User.id)).where(User.role == "owner"))).scalar()
    target = await session.get(User, user_id)
    if target and target.role == "owner" and owners <= 1:
        raise HTTPException(400, "Это последний владелец — удалить нельзя")
    await session.execute(delete(User).where(User.id == user_id))
    log.warning("Удалён пользователь id=%s", user_id)
    return {"ok": True}


# ---------- bots ----------

class BotIn(BaseModel):
    name: str
    token: str | None = None  # None = не менять


def _mask_token(tok: str) -> str:
    return (tok[:6] + "…" + tok[-4:]) if tok and len(tok) > 12 else "…"


@router.get("/bots")
async def list_bots(user=Depends(current_user), session=Depends(get_session)):
    from .bot.runner import manager

    q = select(Bot).order_by(Bot.created_at)
    if user.role != "owner" and user.bot_ids:
        q = q.where(Bot.id.in_([int(x) for x in user.bot_ids]))
    bots = (await session.execute(q)).scalars().all()
    sub_counts = dict(
        (
            await session.execute(
                select(Subscriber.bot_id, func.count(Subscriber.id)).group_by(Subscriber.bot_id)
            )
        ).all()
    )
    funnel_counts = dict(
        (
            await session.execute(
                select(FunnelBot.bot_id, func.count(FunnelBot.id)).group_by(FunnelBot.bot_id)
            )
        ).all()
    )
    return [
        {
            "id": b.id,
            "name": b.name,
            "token_hint": _mask_token(b.token),
            "is_active": b.is_active,
            "running": manager.get(b.id) is not None,
            "tg_username": b.tg_username,
            "last_error": b.last_error,
            "subscribers": sub_counts.get(b.id, 0),
            "funnels": funnel_counts.get(b.id, 0),
        }
        for b in bots
    ]


@router.post("/bots", dependencies=[Depends(require("bots", "edit"))])
async def create_bot(body: BotIn, session=Depends(get_session)):
    if not body.token or not body.token.strip():
        raise HTTPException(400, "Нужен токен бота от @BotFather")
    b = Bot(name=body.name or "Новый бот", token=body.token.strip(), is_active=False)
    session.add(b)
    await session.flush()
    log.info("Добавлен бот «%s» (id=%s)", b.name, b.id)
    return {"id": b.id}


@router.put("/bots/{bot_id}", dependencies=[Depends(require("bots", "edit"))])
async def update_bot(bot_id: int, body: BotIn, session=Depends(get_session)):
    from .bot.runner import manager

    b = await session.get(Bot, bot_id)
    if b is None:
        raise HTTPException(404, "Бот не найден")
    b.name = body.name or b.name
    token_changed = bool(body.token and body.token.strip() and body.token.strip() != b.token)
    if token_changed:
        b.token = body.token.strip()
        b.last_error = None
    await session.commit()
    if token_changed and b.is_active:
        await manager.restart_bot(bot_id)
    return {"ok": True}


@router.post("/bots/{bot_id}/toggle", dependencies=[Depends(require("bots", "edit"))])
async def toggle_bot(bot_id: int, session=Depends(get_session)):
    from .bot.runner import manager

    b = await session.get(Bot, bot_id)
    if b is None:
        raise HTTPException(404, "Бот не найден")
    b.is_active = not b.is_active
    await session.commit()
    if b.is_active:
        await manager.start_bot(bot_id)
    else:
        await manager.stop_bot(bot_id)
    # перечитать возможную ошибку старта
    await session.refresh(b)
    log.info("Бот «%s» (id=%s): %s", b.name, bot_id, "включён" if b.is_active else "выключен")
    if b.is_active and manager.get(bot_id) is None:
        log.error("Бот «%s» не запустился: %s", b.name, b.last_error)
        raise HTTPException(400, f"Бот не запустился: {b.last_error or 'проверьте токен'}")
    return {"is_active": b.is_active, "running": manager.get(bot_id) is not None}


@router.get("/bots/{bot_id}/overview", dependencies=[Depends(require("bots", "view"))])
async def bot_overview(bot_id: int, user=Depends(current_user), session=Depends(get_session)):
    """Сводка для главного экрана бота."""
    from datetime import timedelta

    from .bot.runner import manager

    await ensure_bot_access(user, bot_id)
    b = await session.get(Bot, bot_id)
    if b is None:
        raise HTTPException(404, "Бот не найден")

    subs_q = select(func.count(Subscriber.id)).where(Subscriber.bot_id == bot_id)
    total = (await session.execute(subs_q)).scalar()
    active = (
        await session.execute(subs_q.where(Subscriber.is_active == True))  # noqa: E712
    ).scalar()
    day_ago = datetime.utcnow() - timedelta(hours=24)
    week_ago = datetime.utcnow() - timedelta(days=7)
    new_24h = (
        await session.execute(subs_q.where(Subscriber.created_at >= day_ago))
    ).scalar()
    new_7d = (
        await session.execute(subs_q.where(Subscriber.created_at >= week_ago))
    ).scalar()
    active_24h = (
        await session.execute(subs_q.where(Subscriber.last_active_at >= day_ago))
    ).scalar()
    funnels_n = (
        await session.execute(
            select(func.count(FunnelBot.id)).where(FunnelBot.bot_id == bot_id)
        )
    ).scalar()
    bc_n = (
        await session.execute(
            select(func.count(Broadcast.id)).where(Broadcast.bot_id == bot_id)
        )
    ).scalar()
    return {
        "id": b.id,
        "name": b.name,
        "tg_username": b.tg_username,
        "token_hint": _mask_token(b.token),
        "is_active": b.is_active,
        "running": manager.get(b.id) is not None,
        "last_error": b.last_error,
        "created_at": b.created_at.isoformat(),
        "stats": {
            "subscribers": total,
            "active": active,
            "blocked": total - active,
            "new_24h": new_24h,
            "new_7d": new_7d,
            "active_24h": active_24h,
            "funnels": funnels_n,
            "broadcasts": bc_n,
        },
    }


@router.get("/bots/{bot_id}/check", dependencies=[Depends(require("bots", "view"))])
async def check_bot(bot_id: int, session=Depends(get_session)):
    """Живая диагностика через Telegram API: getMe + getWebhookInfo."""
    from .bot.runner import manager

    b = await session.get(Bot, bot_id)
    if b is None:
        raise HTTPException(404, "Бот не найден")
    from aiogram import Bot as AioBot

    tb = AioBot(token=b.token)
    try:
        me = await tb.get_me()
        wh = await tb.get_webhook_info()
        result = {
            "ok": True,
            "username": me.username,
            "running": manager.get(bot_id) is not None,
            "webhook_url": wh.url or None,
            "pending_updates": wh.pending_update_count,
            "webhook_last_error": wh.last_error_message,
            "hint": (
                "На боте стоит вебхук — long polling не получает сообщения. "
                "Нажмите «Выключить», затем «Включить»: мы снимем вебхук автоматически."
                if wh.url else
                ("Бот не запущен в приложении — нажмите «Включить»." if manager.get(bot_id) is None
                 else "Всё в порядке: вебхука нет, бот запущен.")
            ),
        }
    except Exception as e:  # noqa: BLE001
        result = {"ok": False, "error": str(e)}
    finally:
        await tb.session.close()
    return result


@router.delete("/bots/{bot_id}", dependencies=[Depends(require("bots", "edit")), Depends(require_delete)])
async def delete_bot(bot_id: int, session=Depends(get_session)):
    from .bot.runner import manager

    await manager.stop_bot(bot_id)
    await session.execute(delete(FunnelBot).where(FunnelBot.bot_id == bot_id))
    await session.execute(delete(Bot).where(Bot.id == bot_id))
    log.warning("Удалён бот id=%s", bot_id)
    return {"ok": True}


# ---------- stats ----------

@router.get("/stats", dependencies=[Depends(require("analytics", "view"))])
async def stats(user=Depends(current_user), session=Depends(get_session)):
    mine = await _limit_to_my_bots(user, session, Subscriber.bot_id)
    visible = await _visible_funnel_ids(user, session)

    def sub_q(*extra):
        q = select(func.count(Subscriber.id))
        if mine is not None:
            q = q.where(mine)
        return q.where(*extra) if extra else q

    bots_q = select(func.count(Bot.id))
    my_bots = await allowed_bot_ids(user, session)
    if my_bots is not None:
        bots_q = bots_q.where(Bot.id.in_(my_bots or [-1]))
    bots_n = (await session.execute(bots_q)).scalar()
    subs = (await session.execute(sub_q())).scalar()
    active = (await session.execute(sub_q(Subscriber.is_active == True))).scalar()  # noqa: E712
    fq = select(func.count(Funnel.id))
    if visible is not None:
        fq = fq.where(Funnel.id.in_(visible or [-1]))
    funnels = (await session.execute(fq)).scalar()
    runs = (await session.execute(select(func.count(FunnelRun.id)))).scalar()
    clicks = (await session.execute(select(func.count(ButtonClick.id)))).scalar()
    return {
        "subscribers": subs,
        "active_subscribers": active,
        "funnels": funnels,
        "runs": runs,
        "clicks": clicks,
        "bots": bots_n,
    }


# ---------- аналитика ----------

@router.get("/analytics", dependencies=[Depends(require("analytics", "view"))])
async def analytics(
    days: int = 30,
    bot_id: int | None = None,
    language: str | None = None,
    source: str | None = None,
    tag_id: int | None = None,
    user=Depends(current_user),
    session=Depends(get_session),
):
    from . import analytics as an

    if bot_id:
        await ensure_bot_access(user, bot_id)
    return await an.build_analytics(
        session, days, bot_id, language, source, tag_id,
        allowed_bots=await allowed_bot_ids(user, session),
    )


# ---------- анализ воронок ----------

def _node_label(ntype: str, data: dict) -> str:
    if ntype == "message":
        t = (data.get("text") or "").strip().replace("\n", " ")
        return "💬 " + (t[:40] + "…" if len(t) > 40 else t or "сообщение")
    if ntype == "delay":
        u = {"seconds": "сек", "minutes": "мин", "hours": "ч", "days": "дн"}.get(data.get("unit"), "")
        return f"⏱ задержка {data.get('amount')} {u}"
    if ntype == "condition":
        return "❓ условие"
    if ntype == "action":
        return "🏷 тег"
    if ntype == "start":
        return "▶️ старт"
    return ntype


@router.get("/analysis/options", dependencies=[Depends(require("analytics", "view"))])
async def analysis_options(user=Depends(current_user), session=Depends(get_session)):
    """Справочник для конструктора анализа: воронки с узлами, рассылки."""
    fq = select(Funnel).order_by(Funnel.name)
    visible = await _visible_funnel_ids(user, session)
    if visible is not None:
        fq = fq.where(Funnel.id.in_(visible or [-1]))
    funnels = (await session.execute(fq)).scalars().all()
    fl = []
    for f in funnels:
        nodes = []
        for nid, n in (f.graph.get("nodes") or {}).items():
            if n.get("type") in ("message", "delay", "condition", "action"):
                item = {"id": nid, "label": _node_label(n["type"], n.get("data") or {})}
                if n["type"] == "message":
                    item["buttons"] = [
                        b.get("label", f"кнопка {i+1}")
                        for i, b in enumerate(n.get("data", {}).get("buttons") or [])
                        if not b.get("url")
                    ]
                nodes.append(item)
        fl.append({"id": f.id, "name": f.name, "nodes": nodes})
    bcs = [
        {"id": b.id, "name": b.name}
        for b in (await session.execute(select(Broadcast).order_by(Broadcast.created_at.desc()))).scalars()
    ]
    return {"funnels": fl, "broadcasts": bcs}


class AnalysisIn(BaseModel):
    steps: list
    days: int = 30
    bot_id: int | None = None


@router.post("/analysis/funnel", dependencies=[Depends(require("analytics", "view"))])
async def analysis_funnel(body: AnalysisIn, user=Depends(current_user),
                          session=Depends(get_session)):
    from . import analytics as an

    if not body.steps:
        raise HTTPException(400, "Добавьте хотя бы один шаг")
    if body.bot_id:
        await ensure_bot_access(user, body.bot_id)
    visible = await _visible_funnel_ids(user, session)
    for st in body.steps:
        fid = st.get("funnel_id")
        if fid and visible is not None and int(fid) not in visible:
            raise HTTPException(403, "В шагах есть воронка, к которой нет доступа")
    try:
        return await an.funnel_analysis(
            session, body.steps, body.days, body.bot_id,
            allowed_bots=await allowed_bot_ids(user, session),
        )
    except (KeyError, ValueError, TypeError) as e:
        raise HTTPException(400, f"Некорректный шаг: {e}")


# ---------- tags ----------

class TagIn(BaseModel):
    name: str


@router.get("/tags", dependencies=[Depends(require("tags", "view"))])
async def list_tags(session=Depends(get_session)):
    tags = (await session.execute(select(Tag).order_by(Tag.name))).scalars().all()
    counts = dict(
        (
            await session.execute(
                select(SubscriberTag.tag_id, func.count(SubscriberTag.id)).group_by(
                    SubscriberTag.tag_id
                )
            )
        ).all()
    )
    return [{"id": t.id, "name": t.name, "count": counts.get(t.id, 0)} for t in tags]


@router.post("/tags", dependencies=[Depends(require("tags", "edit"))])
async def create_tag(body: TagIn, session=Depends(get_session)):
    name = body.name.strip()
    if not name:
        raise HTTPException(400, "Пустое имя тега")
    exists = (await session.execute(select(Tag).where(Tag.name == name))).scalar_one_or_none()
    if exists:
        raise HTTPException(400, "Тег уже есть")
    tag = Tag(name=name)
    session.add(tag)
    await session.flush()
    return {"id": tag.id, "name": tag.name}


@router.delete("/tags/{tag_id}", dependencies=[Depends(require("tags", "edit")), Depends(require_delete)])
async def delete_tag(tag_id: int, session=Depends(get_session)):
    await session.execute(delete(SubscriberTag).where(SubscriberTag.tag_id == tag_id))
    await session.execute(delete(Tag).where(Tag.id == tag_id))
    return {"ok": True}


# ---------- subscribers ----------

@router.get("/subscribers", dependencies=[Depends(require("subscribers", "view"))])
async def list_subscribers(
    search: str = "",
    tag_id: int | None = None,
    bot_id: int | None = None,
    user=Depends(current_user),
    session=Depends(get_session),
):
    q = select(Subscriber).order_by(Subscriber.created_at.desc())
    mine = await _limit_to_my_bots(user, session, Subscriber.bot_id)
    if mine is not None:
        q = q.where(mine)
    if bot_id:
        await ensure_bot_access(user, bot_id)
        q = q.where(Subscriber.bot_id == bot_id)
    if search:
        like = f"%{search}%"
        q = q.where(
            Subscriber.username.ilike(like)
            | Subscriber.first_name.ilike(like)
            | Subscriber.last_name.ilike(like)
        )
    if tag_id:
        q = q.where(
            Subscriber.id.in_(
                select(SubscriberTag.subscriber_id).where(SubscriberTag.tag_id == tag_id)
            )
        )
    subs = (await session.execute(q.limit(500))).scalars().all()

    tag_map = await _tags_of(session, [s.id for s in subs])

    return [
        {
            "id": s.id,
            "bot_id": s.bot_id,
            "tg_id": s.tg_id,
            "username": s.username,
            "first_name": s.first_name,
            "last_name": s.last_name,
            "is_active": s.is_active,
            "created_at": s.created_at.isoformat(),
            "tags": tag_map.get(s.id, []),
        }
        for s in subs
    ]


class SubTagIn(BaseModel):
    tag_id: int


@router.post("/subscribers/{sub_id}/tags", dependencies=[Depends(require("subscribers", "edit"))])
async def add_sub_tag(sub_id: int, body: SubTagIn, user=Depends(current_user),
                      session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    sub = await session.get(Subscriber, sub_id)
    if sub is None:
        raise HTTPException(404, "Подписчик не найден")
    exists = (
        await session.execute(
            select(SubscriberTag).where(
                SubscriberTag.subscriber_id == sub_id, SubscriberTag.tag_id == body.tag_id
            )
        )
    ).scalar_one_or_none()
    if exists is None:
        session.add(SubscriberTag(subscriber_id=sub_id, tag_id=body.tag_id))
        await session.flush()
        # триггер воронок "по тегу"
        from .bot.runner import trigger_tag_added

        await trigger_tag_added(session, sub, body.tag_id)
    return {"ok": True}


@router.delete("/subscribers/{sub_id}/tags/{tag_id}", dependencies=[Depends(require("subscribers", "edit"))])
async def remove_sub_tag(sub_id: int, tag_id: int, user=Depends(current_user),
                         session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    await session.execute(
        delete(SubscriberTag).where(
            SubscriberTag.subscriber_id == sub_id, SubscriberTag.tag_id == tag_id
        )
    )
    return {"ok": True}


# ---------- живой чат ----------

@router.get("/subscribers/{sub_id}", dependencies=[Depends(require("subscribers", "view"))])
async def get_subscriber(sub_id: int, user=Depends(current_user),
                         session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    from .bot.runner import manager

    s = await session.get(Subscriber, sub_id)
    if s is None:
        raise HTTPException(404, "Подписчик не найден")
    bot = await session.get(Bot, s.bot_id)
    st = (
        await session.execute(
            select(Tag.id, Tag.name).join(SubscriberTag, SubscriberTag.tag_id == Tag.id)
            .where(SubscriberTag.subscriber_id == sub_id)
        )
    ).all()
    return {
        "id": s.id, "tg_id": s.tg_id,
        "first_name": s.first_name, "last_name": s.last_name, "username": s.username,
        "language_code": s.language_code, "source": s.source,
        "first_source": s.first_source,
        "params": s.params or {},
        "is_active": s.is_active,
        "bot_id": s.bot_id, "bot_name": bot.name if bot else "—",
        "bot_running": manager.get(s.bot_id) is not None,
        "created_at": s.created_at.isoformat(),
        "last_active_at": s.last_active_at.isoformat() if s.last_active_at else None,
        "paused_until": s.automation_paused_until.isoformat() if s.automation_paused_until else None,
        "tags": [{"id": t[0], "name": t[1]} for t in st],
    }


@router.get("/subscribers/{sub_id}/messages", dependencies=[Depends(require("chat", "view"))])
async def get_messages(sub_id: int, user=Depends(current_user),
                       session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    msgs = (
        await session.execute(
            select(Message).where(Message.subscriber_id == sub_id).order_by(Message.created_at)
        )
    ).scalars().all()
    out = [
        {
            "id": m.id, "direction": m.direction, "text": m.text,
            "is_operator": m.is_operator, "is_broadcast": False,
            "created_at": m.created_at.isoformat(),
        }
        for m in msgs
    ]

    # Массовые рассылки в переписку не дублируются: их текст лежит один раз
    # в broadcasts, а факт доставки — в broadcast_recipients. Здесь достаём
    # их и подмешиваем в ленту, чтобы оператор видел полную картину.
    sent = (await session.execute(
        select(BroadcastRecipient.created_at, Broadcast.name, Broadcast.text,
               Broadcast.id, Broadcast.created_at)
        .join(Broadcast, Broadcast.id == BroadcastRecipient.broadcast_id)
        .where(BroadcastRecipient.subscriber_id == sub_id,
               BroadcastRecipient.delivered == True)  # noqa: E712
    )).all()
    for at, name, text, bc_id, bc_at in sent:
        out.append({
            "id": f"bc{bc_id}", "direction": "out", "text": text or "",
            "is_operator": False, "is_broadcast": True, "broadcast_name": name,
            "created_at": (at or bc_at).isoformat(),
        })

    out.sort(key=lambda m: m["created_at"])
    return out


class SendMsgIn(BaseModel):
    text: str = ""
    media: list = []  # [{type, path, name}] — вложения, как в блоке «Сообщение»


@router.post("/subscribers/{sub_id}/send", dependencies=[Depends(require("chat", "edit"))])
async def operator_send(sub_id: int, body: SendMsgIn, user=Depends(current_user),
                        session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    from .bot.runner import manager
    from .bot.sender import send_message_content

    if not body.text.strip() and not body.media:
        raise HTTPException(400, "Пустое сообщение")
    sub = await session.get(Subscriber, sub_id)
    if sub is None:
        raise HTTPException(404, "Подписчик не найден")
    bot = manager.get(sub.bot_id)
    if bot is None:
        raise HTTPException(400, "Бот не запущен — включите его в разделе «Боты»")
    ok = await send_message_content(
        bot, session, sub, body.text, body.media or [], None, is_operator=True
    )
    if not ok:
        raise HTTPException(400, "Не удалось отправить (возможно, подписчик заблокировал бота)")
    return {"ok": True}


class PauseIn(BaseModel):
    minutes: int = 60


@router.post("/subscribers/{sub_id}/pause", dependencies=[Depends(require("chat", "edit"))])
async def pause_automation(sub_id: int, body: PauseIn, user=Depends(current_user),
                           session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    sub = await session.get(Subscriber, sub_id)
    if sub is None:
        raise HTTPException(404, "Подписчик не найден")
    if body.minutes <= 0:
        sub.automation_paused_until = None  # снять паузу
    else:
        sub.automation_paused_until = datetime.utcnow() + timedelta(minutes=body.minutes)
    return {"paused_until": sub.automation_paused_until.isoformat() if sub.automation_paused_until else None}


class StartFlowIn(BaseModel):
    funnel_id: int


@router.post("/subscribers/{sub_id}/start_flow", dependencies=[Depends(require("chat", "edit"))])
async def start_flow(sub_id: int, body: StartFlowIn, user=Depends(current_user),
                     session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    await ensure_funnel_access(user, body.funnel_id)
    from .bot import engine as fx
    from .bot.runner import manager

    sub = await session.get(Subscriber, sub_id)
    funnel = await session.get(Funnel, body.funnel_id)
    if sub is None or funnel is None:
        raise HTTPException(404, "Не найдено")
    if not funnel.graph:
        raise HTTPException(400, "У воронки нет сохранённого графа")
    bot = manager.get(sub.bot_id)
    if bot is None:
        raise HTTPException(400, "Бот не запущен")
    await fx.start_funnel(bot, session, funnel, sub)
    return {"ok": True}


# ---------- segments ----------

@router.get("/segment/fields", dependencies=[Depends(require("subscribers", "view"))])
async def segment_fields(session=Depends(get_session)):
    tags = [
        {"v": t.id, "l": t.name}
        for t in (await session.execute(select(Tag).order_by(Tag.name))).scalars()
    ]
    funnels = [
        {"v": f.id, "l": f.name}
        for f in (await session.execute(select(Funnel).order_by(Funnel.name))).scalars()
    ]
    bcs = [
        {"v": b.id, "l": b.name}
        for b in (await session.execute(select(Broadcast).order_by(Broadcast.created_at.desc()))).scalars()
    ]
    return segment.fields_meta(tags, funnels, bcs)


class SegmentSearchIn(BaseModel):
    bot_id: int | None = None
    filter: dict = {}
    limit: int = 500
    count_only: bool = False


async def _tags_of(session, sub_ids: list[int]) -> dict[int, list]:
    """Теги только для перечисленных подписчиков.

    Раньше здесь выбиралась ВСЯ таблица связок «подписчик — тег» и уже в
    питоне из неё брались нужные пятьсот строк. На пустой базе незаметно, на
    сотне тысяч подписчиков это сотни тысяч строк на КАЖДЫЙ запрос списка —
    и, поскольку приложение однопроцессное, каждый такой запрос вставал
    поперёк всех остальных.
    """
    if not sub_ids:
        return {}
    rows = (await session.execute(
        select(SubscriberTag.subscriber_id, Tag.id, Tag.name)
        .join(Tag, Tag.id == SubscriberTag.tag_id)
        .where(SubscriberTag.subscriber_id.in_(sub_ids))
    )).all()
    out: dict[int, list] = {}
    for sub_id, tid, tname in rows:
        out.setdefault(sub_id, []).append({"id": tid, "name": tname})
    return out


@router.post("/subscribers/search", dependencies=[Depends(require("subscribers", "view"))])
async def subscribers_search(body: SegmentSearchIn, user=Depends(current_user),
                             session=Depends(get_session)):
    if body.bot_id:
        await ensure_bot_access(user, body.bot_id)
    try:
        q = segment.build_query(body.bot_id, body.filter,
                                await allowed_bot_ids(user, session))
    except segment.SegmentError as e:
        raise HTTPException(400, str(e))
    total = (await session.execute(
        select(func.count()).select_from(q.subquery())
    )).scalar()
    if body.count_only:
        return {"total": total, "subscribers": []}

    subs = (await session.execute(
        q.order_by(Subscriber.created_at.desc()).limit(body.limit)
    )).scalars().all()
    tag_map = await _tags_of(session, [s.id for s in subs])
    return {
        "total": total,
        "subscribers": [
            {
                "id": s.id, "bot_id": s.bot_id, "tg_id": s.tg_id,
                "username": s.username, "first_name": s.first_name, "last_name": s.last_name,
                "language_code": s.language_code, "source": s.source,
                "is_active": s.is_active,
                "created_at": s.created_at.isoformat(),
                "last_active_at": s.last_active_at.isoformat() if s.last_active_at else None,
                "tags": tag_map.get(s.id, []),
            }
            for s in subs
        ],
    }


# ---------- удаление и массовые действия с аудиторией ----------

async def _delete_subscribers(session, sub_ids: list[int]) -> int:
    """Удаляет подписчиков со всеми следами: переписка, теги, запуски воронок,
    отложенные задачи, визиты/клики, получатели рассылок."""
    if not sub_ids:
        return 0
    run_ids = select(FunnelRun.id).where(FunnelRun.subscriber_id.in_(sub_ids))
    await session.execute(delete(ScheduledJob).where(ScheduledJob.run_id.in_(run_ids)))
    for model in (Message, SubscriberTag, NodeVisit, ButtonClick,
                  BroadcastRecipient, FunnelRun):
        await session.execute(delete(model).where(model.subscriber_id.in_(sub_ids)))
    res = await session.execute(delete(Subscriber).where(Subscriber.id.in_(sub_ids)))
    return res.rowcount or 0


@router.delete("/subscribers/{sub_id}", dependencies=[Depends(require("subscribers", "edit")), Depends(require_delete)])
async def delete_subscriber(sub_id: int, user=Depends(current_user),
                            session=Depends(get_session)):
    await _my_subscriber(user, session, sub_id)
    sub = (await session.execute(
        select(Subscriber).where(Subscriber.id == sub_id)
    )).scalar_one_or_none()
    if not sub:
        raise HTTPException(404, "Подписчик не найден")
    n = await _delete_subscribers(session, [sub_id])
    log.info("Удалён подписчик id=%s tg_id=%s", sub_id, sub.tg_id)
    return {"ok": True, "deleted": n}


class BulkActionIn(BaseModel):
    bot_id: int | None = None
    filter: dict = {}
    action: str  # delete | add_tag | remove_tag
    tag_id: int | None = None
    expected_total: int | None = None  # защита: сколько человек видел пользователь


@router.post("/subscribers/bulk", dependencies=[Depends(require("subscribers", "edit"))])
async def subscribers_bulk(body: BulkActionIn, user=Depends(current_user),
                           session=Depends(get_session)):
    """Массовое действие над аудиторией, собранной конструктором сегментов."""
    if body.bot_id:
        await ensure_bot_access(user, body.bot_id)
    if body.action == "delete" and not can_delete(user):
        raise HTTPException(403, "Нет права на удаление")
    try:
        q = segment.build_query(body.bot_id, body.filter,
                                await allowed_bot_ids(user, session))
    except segment.SegmentError as e:
        raise HTTPException(400, str(e))

    sub_ids = [s.id for s in (await session.execute(q)).scalars().all()]

    total = len(sub_ids)
    # если аудитория успела измениться с момента показа счётчика — не делаем молча
    if body.expected_total is not None and total != body.expected_total:
        raise HTTPException(409,
            f"Аудитория изменилась: сейчас {total}, вы видели {body.expected_total}. "
            "Обновите список и повторите.")

    if body.action == "delete":
        n = await _delete_subscribers(session, sub_ids)
        log.info("Массовое удаление: %s подписчиков (фильтр %s)", n, body.filter)
        return {"ok": True, "affected": n}

    if body.action in ("add_tag", "remove_tag"):
        if not body.tag_id:
            raise HTTPException(400, "Не указан тег")
        tag = (await session.execute(
            select(Tag).where(Tag.id == body.tag_id)
        )).scalar_one_or_none()
        if not tag:
            raise HTTPException(404, "Тег не найден")
        if body.action == "add_tag":
            have = {r[0] for r in (await session.execute(
                select(SubscriberTag.subscriber_id).where(
                    SubscriberTag.tag_id == body.tag_id,
                    SubscriberTag.subscriber_id.in_(sub_ids),
                )
            )).all()}
            new_ids = [i for i in sub_ids if i not in have]
            session.add_all([
                SubscriberTag(subscriber_id=i, tag_id=body.tag_id) for i in new_ids
            ])
            log.info("Массово добавлен тег «%s»: %s подписчиков", tag.name, len(new_ids))
            return {"ok": True, "affected": len(new_ids)}
        res = await session.execute(delete(SubscriberTag).where(
            SubscriberTag.tag_id == body.tag_id,
            SubscriberTag.subscriber_id.in_(sub_ids),
        ))
        log.info("Массово снят тег «%s»: %s подписчиков", tag.name, res.rowcount or 0)
        return {"ok": True, "affected": res.rowcount or 0}

    raise HTTPException(400, "Неизвестное действие")


# ---------- funnels ----------

class FunnelIn(BaseModel):
    name: str
    trigger_type: str = "start"
    trigger_value: str | None = None
    graph_ui: dict = {}
    bot_ids: list[int] = []


async def _bot_ids_for_funnel(session, funnel_id: int) -> list[int]:
    rows = (
        await session.execute(select(FunnelBot.bot_id).where(FunnelBot.funnel_id == funnel_id))
    ).all()
    return [r[0] for r in rows]


async def _set_funnel_bots(session, funnel_id: int, bot_ids: list[int]):
    await session.execute(delete(FunnelBot).where(FunnelBot.funnel_id == funnel_id))
    for bid in set(bot_ids):
        session.add(FunnelBot(funnel_id=funnel_id, bot_id=bid))


@router.get("/funnels", dependencies=[Depends(require("funnels", "view"))])
async def list_funnels(bot_id: int | None = None, user=Depends(current_user),
                       session=Depends(get_session)):
    q = select(Funnel).order_by(Funnel.created_at.desc())
    visible = await _visible_funnel_ids(user, session)
    if visible is not None:
        q = q.where(Funnel.id.in_(visible or [-1]))
    if bot_id:
        await ensure_bot_access(user, bot_id)
        q = q.join(FunnelBot, FunnelBot.funnel_id == Funnel.id).where(FunnelBot.bot_id == bot_id)
    funnels = (await session.execute(q)).scalars().all()
    run_counts = dict(
        (
            await session.execute(
                select(FunnelRun.funnel_id, func.count(FunnelRun.id)).group_by(
                    FunnelRun.funnel_id
                )
            )
        ).all()
    )
    # карта воронка -> [имена ботов]
    fb = (
        await session.execute(
            select(FunnelBot.funnel_id, Bot.name).join(Bot, Bot.id == FunnelBot.bot_id)
        )
    ).all()
    bots_map: dict[int, list] = {}
    for fid, bname in fb:
        bots_map.setdefault(fid, []).append(bname)
    return [
        {
            "id": f.id,
            "name": f.name,
            "is_active": f.is_active,
            "trigger_type": f.trigger_type,
            "trigger_value": f.trigger_value,
            "runs": run_counts.get(f.id, 0),
            "bots": bots_map.get(f.id, []),
            "updated_at": f.updated_at.isoformat(),
        }
        for f in funnels
    ]


@router.post("/funnels", dependencies=[Depends(require("funnels", "edit"))])
async def create_funnel(body: FunnelIn, user=Depends(current_user),
                        session=Depends(get_session)):
    for bid in body.bot_ids:
        await ensure_bot_access(user, bid)
    funnel = Funnel(
        name=body.name or "Новая воронка",
        trigger_type=body.trigger_type,
        trigger_value=body.trigger_value,
        graph_ui=body.graph_ui or {},
        graph={},
    )
    session.add(funnel)
    await session.flush()
    if body.bot_ids:
        await _set_funnel_bots(session, funnel.id, body.bot_ids)
    return {"id": funnel.id}


@router.get("/funnels/{funnel_id}", dependencies=[Depends(require("funnels", "view"))])
async def get_funnel(funnel_id: int, user=Depends(current_user),
                     session=Depends(get_session)):
    await _my_funnel(user, session, funnel_id)
    f = await session.get(Funnel, funnel_id)
    if f is None:
        raise HTTPException(404, "Воронка не найдена")
    return {
        "id": f.id,
        "name": f.name,
        "is_active": f.is_active,
        "trigger_type": f.trigger_type,
        "trigger_value": f.trigger_value,
        "graph_ui": f.graph_ui,
        "bot_ids": await _bot_ids_for_funnel(session, funnel_id),
    }


@router.put("/funnels/{funnel_id}", dependencies=[Depends(require("funnels", "edit"))])
async def update_funnel(funnel_id: int, body: FunnelIn, user=Depends(current_user),
                        session=Depends(get_session)):
    await _my_funnel(user, session, funnel_id)
    for bid in body.bot_ids:
        await ensure_bot_access(user, bid)
    f = await session.get(Funnel, funnel_id)
    if f is None:
        raise HTTPException(404, "Воронка не найдена")
    try:
        compiled = compile_graph(body.graph_ui)
    except GraphError as e:
        raise HTTPException(400, str(e))
    f.name = body.name
    f.trigger_type = body.trigger_type
    f.trigger_value = body.trigger_value
    f.graph_ui = body.graph_ui
    f.graph = compiled
    f.updated_at = datetime.utcnow()
    await _set_funnel_bots(session, funnel_id, body.bot_ids)
    return {"ok": True}


@router.post("/funnels/{funnel_id}/toggle", dependencies=[Depends(require("funnels", "edit"))])
async def toggle_funnel(funnel_id: int, user=Depends(current_user),
                        session=Depends(get_session)):
    await _my_funnel(user, session, funnel_id)
    f = await session.get(Funnel, funnel_id)
    if f is None:
        raise HTTPException(404, "Воронка не найдена")
    if not f.is_active and not f.graph:
        raise HTTPException(400, "Сначала сохраните воронку с корректным графом")
    if not f.is_active and not await _bot_ids_for_funnel(session, funnel_id):
        raise HTTPException(400, "Назначьте воронке хотя бы одного бота")
    f.is_active = not f.is_active
    log.info("Воронка «%s» (id=%s): %s", f.name, funnel_id,
             "включена" if f.is_active else "выключена")
    return {"is_active": f.is_active}


@router.delete("/funnels/{funnel_id}", dependencies=[Depends(require("funnels", "edit")), Depends(require_delete)])
async def delete_funnel(funnel_id: int, user=Depends(current_user),
                        session=Depends(get_session)):
    await _my_funnel(user, session, funnel_id)
    await session.execute(delete(Funnel).where(Funnel.id == funnel_id))
    return {"ok": True}


@router.get("/funnels/{funnel_id}/stats", dependencies=[Depends(require("funnels", "view"))])
async def funnel_stats(funnel_id: int, user=Depends(current_user),
                       session=Depends(get_session)):
    await _my_funnel(user, session, funnel_id)
    from .models import NodeVisit

    total = (
        await session.execute(
            select(func.count(FunnelRun.id)).where(FunnelRun.funnel_id == funnel_id)
        )
    ).scalar()
    uniq_entered = (
        await session.execute(
            select(func.count(func.distinct(FunnelRun.subscriber_id))).where(
                FunnelRun.funnel_id == funnel_id
            )
        )
    ).scalar()
    done = (
        await session.execute(
            select(func.count(FunnelRun.id)).where(
                FunnelRun.funnel_id == funnel_id, FunnelRun.status == "done"
            )
        )
    ).scalar()
    # уникальные подписчики, дошедшие до каждого узла
    visits = (
        await session.execute(
            select(NodeVisit.node_id, func.count(func.distinct(NodeVisit.subscriber_id)))
            .where(NodeVisit.funnel_id == funnel_id)
            .group_by(NodeVisit.node_id)
        )
    ).all()
    clicks = (
        await session.execute(
            select(
                ButtonClick.node_id, ButtonClick.button_index,
                func.count(func.distinct(ButtonClick.subscriber_id)),
            )
            .where(ButtonClick.funnel_id == funnel_id)
            .group_by(ButtonClick.node_id, ButtonClick.button_index)
        )
    ).all()
    return {
        "runs": total,
        "unique_entered": uniq_entered,
        "done": done,
        "nodes": {n: c for n, c in visits},
        "clicks": [
            {"node_id": n, "button": b, "count": c} for n, b, c in clicks
        ],
    }


# ---------- интеграция с Google Таблицами ----------

SHEETS_KEY = "sheets_integration"

DEFAULT_SHEETS_CFG = {
    "spreadsheet_id": "",
    "credentials": {},          # ключ сервисного аккаунта
    "auto": False,              # выгружать по расписанию
    "interval": "daily",        # daily | hourly
    "hour": 4,                  # час ночной выгрузки (UTC), если daily
    "days": 30,                 # период для отчётов «за N дней»
    "sheets": {"funnels": True},
    "last_run": None,
    "last_status": None,        # ok | error
    "last_error": "",
    "last_counts": {},
}


async def get_sheets_cfg(session) -> dict:
    from .models import Setting

    row = (await session.execute(
        select(Setting).where(Setting.key == SHEETS_KEY))).scalar_one_or_none()
    cfg = dict(DEFAULT_SHEETS_CFG)
    cfg.update(row.value or {} if row else {})
    return cfg


async def save_sheets_cfg(session, cfg: dict):
    from .models import Setting

    row = (await session.execute(
        select(Setting).where(Setting.key == SHEETS_KEY))).scalar_one_or_none()
    if row:
        row.value = cfg
        flag_modified(row, "value")
    else:
        session.add(Setting(key=SHEETS_KEY, value=cfg))
    await session.flush()


def _sheets_public(cfg: dict) -> dict:
    """Наружу отдаём всё, кроме приватного ключа."""
    cred = cfg.get("credentials") or {}
    return {
        "spreadsheet_id": cfg.get("spreadsheet_id", ""),
        "connected": bool(cred.get("client_email")),
        "robot_email": cred.get("client_email"),
        "project_id": cred.get("project_id"),
        "auto": bool(cfg.get("auto")),
        "interval": cfg.get("interval", "daily"),
        "hour": cfg.get("hour", 4),
        "days": cfg.get("days", 30),
        "sheets": cfg.get("sheets") or {"funnels": True},
        "last_run": cfg.get("last_run"),
        "last_status": cfg.get("last_status"),
        "last_error": cfg.get("last_error") or "",
        "last_counts": cfg.get("last_counts") or {},
        "available": exports.SHEETS,
    }


@router.get("/integrations/sheets", dependencies=[Depends(require("integrations", "view"))])
async def sheets_get(session=Depends(get_session)):
    return _sheets_public(await get_sheets_cfg(session))


class SheetsCfgIn(BaseModel):
    spreadsheet_id: str = ""
    credentials_json: str | None = None   # None = не менять
    auto: bool = False
    interval: str = "daily"
    hour: int = 4
    days: int = 30
    sheets: dict = {"funnels": True}


@router.put("/integrations/sheets", dependencies=[Depends(require("integrations", "edit"))])
async def sheets_save(body: SheetsCfgIn, session=Depends(get_session)):
    cfg = await get_sheets_cfg(session)

    if body.credentials_json:
        try:
            cred = json.loads(body.credentials_json)
        except Exception:  # noqa: BLE001
            raise HTTPException(400, "Это не похоже на JSON. Вставьте содержимое "
                                     "файла-ключа целиком, вместе со скобками.")
        if not cred.get("client_email") or not cred.get("private_key"):
            raise HTTPException(400, "В файле нет client_email или private_key — "
                                     "похоже, это не ключ сервисного аккаунта.")
        cfg["credentials"] = cred

    # из ссылки на таблицу вытащим id сами — так проще, чем объяснять, что копировать
    sid = (body.spreadsheet_id or "").strip()
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", sid)
    if m:
        sid = m.group(1)
    cfg["spreadsheet_id"] = sid

    cfg["auto"] = bool(body.auto)
    cfg["interval"] = body.interval if body.interval in ("daily", "hourly") else "daily"
    cfg["hour"] = max(0, min(23, int(body.hour)))
    cfg["days"] = max(2, min(365, int(body.days)))
    chosen = {k: bool(v) for k, v in (body.sheets or {}).items() if k in exports.SHEET_KEYS}
    if not any(chosen.values()):
        chosen = {"funnels": True}
    cfg["sheets"] = chosen
    await save_sheets_cfg(session, cfg)
    log.info("Настройки выгрузки в Google Таблицы обновлены (авто: %s)", cfg["auto"])
    return _sheets_public(cfg)


async def run_sheets_export(session, cfg: dict, allowed_bots=None) -> dict:
    """Собирает выбранные листы и пишет их в таблицу. Возвращает {лист: строк}."""
    from . import sheets as gs

    if not (cfg.get("credentials") or {}).get("client_email"):
        raise gs.SheetsError("Сначала вставьте ключ сервисного аккаунта")
    tabs = await exports.build_tabs(
        session, cfg.get("sheets") or {}, cfg.get("days", 30), allowed_bots)
    return await gs.write_tabs(cfg["credentials"], cfg["spreadsheet_id"], tabs)


@router.post("/integrations/sheets/export",
             dependencies=[Depends(require("integrations", "edit"))])
async def sheets_export_now(user=Depends(current_user), session=Depends(get_session)):
    from . import sheets as gs

    cfg = await get_sheets_cfg(session)
    try:
        counts = await run_sheets_export(session, cfg, await allowed_bot_ids(user, session))
    except gs.SheetsError as e:
        cfg["last_run"] = datetime.utcnow().isoformat()
        cfg["last_status"] = "error"
        cfg["last_error"] = str(e)
        await save_sheets_cfg(session, cfg)
        log.warning("Выгрузка в Google Таблицы не удалась: %s", e)
        raise HTTPException(400, str(e))

    cfg["last_run"] = datetime.utcnow().isoformat()
    cfg["last_status"] = "ok"
    cfg["last_error"] = ""
    cfg["last_counts"] = counts
    await save_sheets_cfg(session, cfg)
    log.info("Выгрузка в Google Таблицы: %s", counts)
    return {"ok": True, "counts": counts,
            "url": f"https://docs.google.com/spreadsheets/d/{cfg['spreadsheet_id']}"}


@router.get("/integrations/sheets/preview",
            dependencies=[Depends(require("integrations", "view"))])
async def sheets_preview(user=Depends(current_user), session=Depends(get_session)):
    """Первые строки того, что уйдёт в таблицу — посмотреть перед выгрузкой."""
    cfg = await get_sheets_cfg(session)
    tabs = await exports.build_tabs(
        session, cfg.get("sheets") or {}, cfg.get("days", 30),
        await allowed_bot_ids(user, session))
    return {t: {"rows": len(rows) - 1, "sample": rows[:6]} for t, rows in tabs.items()}


# ---------- AI ----------

class AISettingsIn(BaseModel):
    provider: str = "anthropic"
    model: str | None = None
    api_key: str | None = None  # None = не менять


@router.get("/ai/settings", dependencies=[Depends(require("ai", "view"))])
async def ai_settings(session=Depends(get_session)):
    s = await ai.get_ai_settings(session)
    key = s.get("api_key") or ""
    return {
        "provider": s.get("provider", "anthropic"),
        "model": s.get("model") or ai.DEFAULT_MODELS.get(s.get("provider", "anthropic")),
        "has_key": bool(key),
        "key_hint": ("…" + key[-4:]) if key else None,
    }


@router.put("/ai/settings", dependencies=[Depends(require_owner)])
async def ai_settings_save(body: AISettingsIn, session=Depends(get_session)):
    if body.provider not in ("anthropic", "openai"):
        raise HTTPException(400, "Провайдер: anthropic или openai")
    s = await ai.get_ai_settings(session)
    s["provider"] = body.provider
    s["model"] = body.model or ai.DEFAULT_MODELS[body.provider]
    if body.api_key:
        s["api_key"] = body.api_key.strip()
    await ai.save_ai_settings(session, s)
    return {"ok": True}


@router.get("/ai/usage", dependencies=[Depends(require("ai", "view"))])
async def ai_usage(session=Depends(get_session)):
    rows = (
        (await session.execute(select(AIRequest).order_by(AIRequest.created_at.desc()).limit(20)))
        .scalars().all()
    )
    tot_in = (await session.execute(select(func.sum(AIRequest.input_tokens)))).scalar() or 0
    tot_out = (await session.execute(select(func.sum(AIRequest.output_tokens)))).scalar() or 0
    tot_n = (await session.execute(select(func.count(AIRequest.id)))).scalar() or 0
    return {
        "requests": tot_n,
        "input_tokens": tot_in,
        "output_tokens": tot_out,
        "recent": [
            {
                "id": r.id, "provider": r.provider, "model": r.model,
                "status": r.status, "error": r.error,
                "input_tokens": r.input_tokens, "output_tokens": r.output_tokens,
                "funnel_id": r.funnel_id, "created_at": r.created_at.isoformat(),
            }
            for r in rows
        ],
    }


def _guess_kind(content_type: str, ext: str) -> str:
    ct = content_type.lower()
    if ct.startswith("image/"):
        return "photo"
    if ct.startswith("video/"):
        return "video"
    if ct.startswith("audio/"):
        return "voice" if ext in (".ogg", ".oga") else "audio"
    if ext in (".mp4", ".mov", ".webm", ".mkv"):
        return "video"
    if ext in (".mp3", ".m4a", ".wav", ".flac"):
        return "audio"
    if ext in (".ogg", ".oga"):
        return "voice"
    if ext in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        return "photo"
    return "document"


async def require_any_editor(user=Depends(current_user)):
    """Загрузка файлов: нужна хотя бы одна возможность что-то редактировать."""
    from .auth import has_perm
    if any(has_perm(user, f, "edit") for f in ("funnels", "broadcasts", "chat")):
        return user
    raise HTTPException(403, "Нет прав на загрузку файлов")


@router.post("/media/upload", dependencies=[Depends(require_any_editor)])
async def upload_media(file: UploadFile):
    """Загрузка любого вложения для блока «Сообщение».
    Возвращает {path, kind, name}. kind — предполагаемый тип (пользователь может сменить)."""
    import uuid
    from pathlib import Path

    from .config import settings

    data = await file.read()
    if len(data) > 50 * 1024 * 1024:  # лимит бота Telegram — 50 МБ
        raise HTTPException(400, "Слишком большой файл (лимит 50 МБ для ботов Telegram)")
    orig = file.filename or "file"
    ext = Path(orig).suffix.lower()
    kind = _guess_kind(file.content_type or "", ext)
    media = Path(settings.media_dir)
    media.mkdir(parents=True, exist_ok=True)
    fname = f"f_{uuid.uuid4().hex[:10]}{ext or ''}"
    (media / fname).write_bytes(data)
    return {"path": f"media/{fname}", "kind": kind, "name": orig}


@router.post("/ai/extract_docx", dependencies=[Depends(require("ai", "edit"))])
async def extract_docx(file: UploadFile):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(400, "Нужен файл .docx")
    import io
    import uuid
    from pathlib import Path

    from .config import settings

    try:
        from docx import Document

        doc = Document(io.BytesIO(await file.read()))
    except Exception:  # noqa: BLE001
        raise HTTPException(400, "Не удалось прочитать файл — это точно .docx?")

    media = Path(settings.media_dir)
    media.mkdir(parents=True, exist_ok=True)
    saved_images = 0

    def para_text(p) -> str:
        """Текст абзаца + маркеры картинок на месте их вставки."""
        nonlocal saved_images
        out = []
        for run in p.runs:
            blips = run.element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for blip in blips:
                rid = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                part = p.part.related_parts.get(rid) if rid else None
                if part is None:
                    continue
                ext = Path(str(part.partname)).suffix or ".png"
                fname = f"tz_{uuid.uuid4().hex[:8]}{ext}"
                (media / fname).write_bytes(part.blob)
                saved_images += 1
                out.append(f"[КАРТИНКА: media/{fname}]")
            if run.text:
                out.append(run.text)
        return "".join(out)

    parts: list[str] = []
    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        for item in doc.iter_inner_content():
            if isinstance(item, Paragraph):
                t = para_text(item)
                if t.strip():
                    style = (item.style.name or "").lower() if item.style else ""
                    parts.append(("## " + t) if style.startswith("heading") else t)
            elif isinstance(item, Table):
                for row in item.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
    except AttributeError:  # старый python-docx без iter_inner_content
        parts = [para_text(p) for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))

    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(400, "В документе не нашлось текста")
    return {"text": text, "chars": len(text), "images": saved_images}


class GenerateIn(BaseModel):
    spec_text: str


@router.post("/ai/generate", dependencies=[Depends(require("ai", "edit"))])
async def ai_generate(body: GenerateIn, session=Depends(get_session)):
    if not body.spec_text.strip():
        raise HTTPException(400, "Пустое ТЗ")
    s = await ai.get_ai_settings(session)
    if not s.get("api_key"):
        raise HTTPException(400, "Сначала укажите API-ключ в настройках AI")
    provider = s.get("provider", "anthropic")
    model = s.get("model") or ai.DEFAULT_MODELS[provider]

    req = AIRequest(provider=provider, model=model)
    session.add(req)
    await session.flush()
    try:
        text, tin, tout = await ai.call_llm(provider, s["api_key"], model, body.spec_text)
        req.input_tokens, req.output_tokens = tin, tout
        spec = ai.parse_llm_json(text)
        tag_ids = await ai.ensure_tags(session, spec)
        fields = ai.spec_to_funnel_fields(spec, tag_ids)
    except ai.AIError as e:
        req.status, req.error = "error", str(e)
        await session.commit()  # сохранить запись об ошибке в статистику
        raise HTTPException(400, str(e))
    except Exception as e:  # noqa: BLE001
        req.status, req.error = "error", str(e)
        await session.commit()
        raise HTTPException(500, f"Ошибка генерации: {e}")

    funnel = Funnel(is_active=False, **fields)
    session.add(funnel)
    await session.flush()
    req.funnel_id = funnel.id
    return {"funnel_id": funnel.id, "name": funnel.name,
            "input_tokens": req.input_tokens, "output_tokens": req.output_tokens}


# ---------- AI-чат воронки ----------

class AIEditIn(BaseModel):
    funnel_id: int
    messages: list  # [{role: user|assistant, content: str}]


@router.post("/ai/edit", dependencies=[Depends(require("ai", "edit"))])
async def ai_edit(body: AIEditIn, session=Depends(get_session)):
    s = await ai.get_ai_settings(session)
    if not s.get("api_key"):
        raise HTTPException(400, "Сначала укажите API-ключ в разделе «AI-сборка»")
    funnel = await session.get(Funnel, body.funnel_id)
    if funnel is None:
        raise HTTPException(404, "Воронка не найдена")
    if not body.messages:
        raise HTTPException(400, "Пустой запрос")
    provider = s.get("provider", "anthropic")
    model = s.get("model") or ai.DEFAULT_MODELS[provider]
    tags_list = (await session.execute(select(Tag))).scalars().all()

    req = AIRequest(provider=provider, model=model)
    session.add(req)
    await session.flush()
    try:
        reply, fields, tin, tout = await ai.chat_edit_funnel(
            session, funnel, tags_list, body.messages, provider, s["api_key"], model
        )
        req.input_tokens, req.output_tokens = tin, tout
        req.funnel_id = funnel.id
    except ai.AIError as e:
        req.status, req.error = "error", str(e)
        await session.commit()
        raise HTTPException(400, str(e))
    except Exception as e:  # noqa: BLE001
        req.status, req.error = "error", str(e)
        await session.commit()
        raise HTTPException(500, f"Ошибка AI: {e}")

    updated = False
    if fields:
        funnel.name = fields["name"]
        funnel.trigger_type = fields["trigger_type"]
        funnel.trigger_value = fields["trigger_value"]
        funnel.graph_ui = fields["graph_ui"]
        funnel.graph = fields["graph"]
        funnel.updated_at = datetime.utcnow()
        updated = True
    # фиксируем сразу и отдаём готовый граф — редактор рисует его без второго запроса
    await session.commit()
    return {
        "reply": reply,
        "updated": updated,
        "tokens": req.input_tokens + req.output_tokens,
        "funnel": {
            "id": funnel.id,
            "name": funnel.name,
            "trigger_type": funnel.trigger_type,
            "trigger_value": funnel.trigger_value,
            "graph_ui": funnel.graph_ui,
        } if updated else None,
    }


# ---------- broadcasts ----------

class BroadcastIn(BaseModel):
    name: str
    text: str
    bot_id: int
    photo_url: str | None = None
    media: list = []  # [{type, path, name}] — вложения (в т.ч. видео/альбомы)
    include_tags: list[int] = []
    exclude_tags: list[int] = []
    segment: dict | None = None  # если задан — используется вместо include/exclude


@router.get("/broadcasts", dependencies=[Depends(require("broadcasts", "view"))])
async def list_broadcasts(bot_id: int | None = None, user=Depends(current_user),
                          session=Depends(get_session)):
    q = select(Broadcast).order_by(Broadcast.created_at.desc())
    mine = await _limit_to_my_bots(user, session, Broadcast.bot_id)
    if mine is not None:
        q = q.where(mine)
    if bot_id:
        await ensure_bot_access(user, bot_id)
        q = q.where(Broadcast.bot_id == bot_id)
    bcs = (await session.execute(q)).scalars().all()
    bot_names = dict((await session.execute(select(Bot.id, Bot.name))).all())
    return [
        {
            "id": b.id,
            "name": b.name,
            "text": b.text,
            "bot": bot_names.get(b.bot_id, "—"),
            "status": b.status,
            "total": b.total,
            "sent": b.sent,
            "failed": b.failed,
            "created_at": b.created_at.isoformat(),
        }
        for b in bcs
    ]


@router.get("/broadcasts/{bc_id}", dependencies=[Depends(require("broadcasts", "view"))])
async def broadcast_detail(bc_id: int, user=Depends(current_user),
                           session=Depends(get_session)):
    """Карточка рассылки: что отправляли, кому и с каким результатом."""
    bc = await session.get(Broadcast, bc_id)
    if not bc:
        raise HTTPException(404, "Рассылка не найдена")
    await ensure_bot_access(user, bc.bot_id)

    bot = await session.get(Bot, bc.bot_id)

    # справочники для расшифровки условий сегмента
    tag_names = dict((await session.execute(select(Tag.id, Tag.name))).all())
    funnel_names = dict((await session.execute(select(Funnel.id, Funnel.name))).all())
    bc_names = dict((await session.execute(select(Broadcast.id, Broadcast.name))).all())

    f = bc.filters or {}
    if "segment" in f:
        audience = segment.describe(f.get("segment") or {}, tag_names, funnel_names, bc_names)
        audience_kind = "Сегмент"
    else:
        audience_kind = "Теги"
        audience = []
        inc = [tag_names.get(t, f"#{t}") for t in (f.get("include_tags") or [])]
        exc = [tag_names.get(t, f"#{t}") for t in (f.get("exclude_tags") or [])]
        if inc:
            audience.append("есть тег: " + ", ".join(inc))
        if exc:
            audience.append("нет тега: " + ", ".join(exc))
    if not audience:
        audience = ["вся база бота"]

    # фактическая доставка
    delivered, not_delivered = (await session.execute(
        select(
            func.sum(case((BroadcastRecipient.delivered == True, 1), else_=0)),  # noqa: E712
            func.sum(case((BroadcastRecipient.delivered == False, 1), else_=0)),  # noqa: E712
        ).where(BroadcastRecipient.broadcast_id == bc_id)
    )).one()

    # первые 100 получателей — чтобы можно было глазами проверить, кому ушло
    rows = (await session.execute(
        select(Subscriber, BroadcastRecipient.delivered, BroadcastRecipient.created_at)
        .join(BroadcastRecipient, BroadcastRecipient.subscriber_id == Subscriber.id)
        .where(BroadcastRecipient.broadcast_id == bc_id)
        .order_by(BroadcastRecipient.created_at.desc())
        .limit(100)
    )).all()

    return {
        "id": bc.id,
        "name": bc.name,
        "text": bc.text,
        "media": bc.media or [],
        "photo_url": bc.photo_url,
        "bot": bot.name if bot else "—",
        "bot_id": bc.bot_id,
        "status": bc.status,
        "total": bc.total,
        "sent": bc.sent,
        "failed": bc.failed,
        "delivered": int(delivered or 0),
        "not_delivered": int(not_delivered or 0),
        "created_at": bc.created_at.isoformat(),
        "audience_kind": audience_kind,
        "audience": audience,
        "recipients": [
            {
                "id": s.id,
                "name": (f"{s.first_name or ''} {s.last_name or ''}").strip() or "—",
                "username": s.username,
                "delivered": bool(d),
                "at": (at or bc.created_at).isoformat(),
            }
            for s, d, at in rows
        ],
    }


@router.post("/broadcasts", dependencies=[Depends(require("broadcasts", "edit"))])
async def create_broadcast(body: BroadcastIn, user=Depends(current_user),
                           session=Depends(get_session)):
    await ensure_bot_access(user, body.bot_id)
    if not body.text.strip() and not body.media:
        raise HTTPException(400, "Добавьте текст или вложение")
    if not await session.get(Bot, body.bot_id):
        raise HTTPException(400, "Выберите бота для рассылки")
    if body.segment is not None:
        filters = {"segment": body.segment}
    else:
        filters = {"include_tags": body.include_tags, "exclude_tags": body.exclude_tags}
    bc = Broadcast(
        bot_id=body.bot_id,
        name=body.name or "Рассылка",
        text=body.text,
        photo_url=body.photo_url,
        media=body.media or [],
        filters=filters,
    )
    session.add(bc)
    await session.flush()
    log.info("Создана рассылка «%s» (id=%s, бот %s, вложений %s)",
             bc.name, bc.id, bc.bot_id, len(bc.media or []))
    return {"id": bc.id}


# ---------- конверсии по шагам в лист вида 08B_FUNNEL_INPUT ----------
# Отдельная от «обычной» выгрузки история: там листы наши и мы их затираем,
# здесь лист чужой, в нём есть ручные колонки и формулы — пишем точечно.
# Ключ робота общий, берём из настроек выше.

FI_KEY = "funnel_input_integration"

DEFAULT_FI_CFG = {
    "spreadsheet_id": "",
    "sheet_name": "08B_FUNNEL_INPUT",
    "auto": False,
    "interval": "daily",        # weekly | daily | hourly
    "hour": 4,                  # час выгрузки (UTC) для daily и weekly
    "start_week": None,         # с какой недели начали писать
    "last_run": None,
    "last_status": None,
    "last_error": "",
    "last_result": {},
}


async def get_fi_cfg(session) -> dict:
    from .models import Setting

    row = (await session.execute(
        select(Setting).where(Setting.key == FI_KEY))).scalar_one_or_none()
    cfg = dict(DEFAULT_FI_CFG)
    cfg.update(row.value or {} if row else {})
    return cfg


async def save_fi_cfg(session, cfg: dict):
    from .models import Setting

    row = (await session.execute(
        select(Setting).where(Setting.key == FI_KEY))).scalar_one_or_none()
    if row:
        row.value = cfg
        flag_modified(row, "value")
    else:
        session.add(Setting(key=FI_KEY, value=cfg))
    await session.flush()


def _fi_public(cfg: dict) -> dict:
    return {
        "spreadsheet_id": cfg.get("spreadsheet_id", ""),
        "sheet_name": cfg.get("sheet_name") or "08B_FUNNEL_INPUT",
        "auto": bool(cfg.get("auto")),
        "interval": cfg.get("interval", "daily"),
        "hour": cfg.get("hour", 4),
        "start_week": cfg.get("start_week"),
        "last_run": cfg.get("last_run"),
        "last_status": cfg.get("last_status"),
        "last_error": cfg.get("last_error") or "",
        "last_result": cfg.get("last_result") or {},
    }


@router.get("/integrations/funnel-input",
            dependencies=[Depends(require("integrations", "view"))])
async def fi_get(session=Depends(get_session)):
    cfg = await get_fi_cfg(session)
    sheets_cfg = await get_sheets_cfg(session)
    cred = sheets_cfg.get("credentials") or {}
    out = _fi_public(cfg)
    out["connected"] = bool(cred.get("client_email"))
    out["robot_email"] = cred.get("client_email")
    return out


class FunnelInputCfgIn(BaseModel):
    spreadsheet_id: str = ""
    sheet_name: str = "08B_FUNNEL_INPUT"
    auto: bool = False
    interval: str = "daily"
    hour: int = 4


@router.put("/integrations/funnel-input",
            dependencies=[Depends(require("integrations", "edit"))])
async def fi_save(body: FunnelInputCfgIn, session=Depends(get_session)):
    cfg = await get_fi_cfg(session)

    sid = (body.spreadsheet_id or "").strip()
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", sid)
    if m:
        sid = m.group(1)
    cfg["spreadsheet_id"] = sid
    cfg["sheet_name"] = (body.sheet_name or "").strip() or "08B_FUNNEL_INPUT"
    cfg["auto"] = bool(body.auto)
    cfg["interval"] = (body.interval
                       if body.interval in ("weekly", "daily", "hourly") else "daily")
    cfg["hour"] = max(0, min(23, int(body.hour)))
    await save_fi_cfg(session, cfg)
    log.info("Настройки выгрузки конверсий обновлены (авто: %s, %s)",
             cfg["auto"], cfg["interval"])
    return _fi_public(cfg)


async def run_funnel_input_export(session, cfg: dict, allowed_bots=None) -> dict:
    """Собирает недельные конверсии и дописывает их в лист."""
    from . import funnel_input as fi
    from . import sheets as gs

    cred = (await get_sheets_cfg(session)).get("credentials") or {}
    if not cred.get("client_email"):
        raise gs.SheetsError(
            "Сначала подключите робота в блоке выше — вставьте ключ сервисного аккаунта")

    weeks = fi.weeks_since(cfg.get("start_week"))
    rows = await fi.collect_rows(session, weeks, allowed_bots)
    result = await fi.write_rows(
        cred, cfg.get("spreadsheet_id", ""),
        cfg.get("sheet_name") or "08B_FUNNEL_INPUT", rows)
    result["weeks"] = len(weeks)
    return result


@router.post("/integrations/funnel-input/export",
             dependencies=[Depends(require("integrations", "edit"))])
async def fi_export_now(user=Depends(current_user), session=Depends(get_session)):
    from . import funnel_input as fi
    from . import sheets as gs

    cfg = await get_fi_cfg(session)
    try:
        result = await run_funnel_input_export(
            session, cfg, await allowed_bot_ids(user, session))
    except gs.SheetsError as e:
        cfg["last_run"] = datetime.utcnow().isoformat()
        cfg["last_status"] = "error"
        cfg["last_error"] = str(e)
        await save_fi_cfg(session, cfg)
        log.warning("Выгрузка конверсий не удалась: %s", e)
        raise HTTPException(400, str(e))

    if not cfg.get("start_week"):
        cfg["start_week"] = fi.monday(datetime.utcnow().date()).isoformat()
    cfg["last_run"] = datetime.utcnow().isoformat()
    cfg["last_status"] = "ok"
    cfg["last_error"] = ""
    cfg["last_result"] = result
    await save_fi_cfg(session, cfg)
    log.info("Выгрузка конверсий: %s", result)
    return {"ok": True, "result": result,
            "url": f"https://docs.google.com/spreadsheets/d/{cfg['spreadsheet_id']}"}


@router.get("/integrations/funnel-input/preview",
            dependencies=[Depends(require("integrations", "view"))])
async def fi_preview(user=Depends(current_user), session=Depends(get_session)):
    """Что уйдёт в лист — посмотреть, не трогая таблицу."""
    from . import funnel_input as fi

    cfg = await get_fi_cfg(session)
    weeks = fi.weeks_since(cfg.get("start_week"))
    rows = await fi.collect_rows(session, weeks, await allowed_bot_ids(user, session))
    return {
        "weeks": [w[0].strftime("%Y-%m-%d") for w in weeks],
        "rows": len(rows),
        "sample": [
            {"week": r["week"], "label": r["label"],
             "steps": r["steps"][:12], "total_steps": len(r["steps"])}
            for r in rows[:12]
        ],
    }
